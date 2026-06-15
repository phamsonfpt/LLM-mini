import os
import csv
from docx import Document
from fpdf import FPDF
from PIL import Image, ImageDraw, ImageFont
from gtts import gTTS

def ensure_dir(path):
    if not os.path.exists(path):
        os.makedirs(path)

def generate_complex_test_data():
    output_dir = r"d:\LLM_mini\test_data"
    ensure_dir(output_dir)
    print(f"Generating complex interconnected test files in {output_dir}...")

    # 1. Markdown: Nhiệm vụ
    md_content = """# Hồ sơ Chiến dịch Cứu hộ K-9
Sứ mệnh giải cứu trạm không gian K-9 đang xoay quanh 3 con tàu: Apollo, Hermes, và Zeus.
- **Tàu Apollo**: Chở hàng hóa y tế.
- **Tàu Hermes**: Trạm chỉ huy di động, đang hoạt động ổn định.
- **Tàu Zeus**: Tàu thăm dò lõi, có lịch sử bị lỗi phần cứng thường xuyên.

Nếu bất kỳ tàu nào gặp sự cố rò rỉ khí nghiêm trọng, phải lập tức kích hoạt Giao thức 77-A để phong tỏa khu vực. Tuyệt đối không được mở cửa khoang tự do.
"""
    with open(os.path.join(output_dir, "01_SoTayNhiemVu.md"), "w", encoding="utf-8") as f:
        f.write(md_content)

    # 2. Text: Nhân sự
    txt_content = """HỒ SƠ NHÂN SỰ CHIẾN DỊCH K-9:
- Chỉ huy trưởng: Đại tá Elena. Hiện đang ở Trạm Trái Đất.
- Kỹ sư trưởng: Tiến sĩ Aris. Anh là người thiết kế hệ thống van oxy cho cả 3 con tàu. Hiện tại, Tiến sĩ Aris đang được phân công đóng quân trên tàu Hermes để giám sát toàn bộ chiến dịch.
- Phi công chính: Mark Watney, hiện đang cầm lái con tàu Zeus.
Lưu ý: Chỉ duy nhất Kỹ sư trưởng mới có thẩm quyền cao nhất trong việc thay thế các bộ phận lõi của hệ thống sinh tồn.
"""
    with open(os.path.join(output_dir, "02_HoSoNhanSu.txt"), "w", encoding="utf-8") as f:
        f.write(txt_content)

    # 3. DOCX: Quy định
    doc = Document()
    doc.add_heading('Sổ tay Quy định Nội bộ', 0)
    doc.add_paragraph('Chi tiết Giao thức 77-A (Xử lý rò rỉ khí khẩn cấp):')
    doc.add_paragraph('1. Khi phát hiện rò rỉ, toàn bộ phi hành đoàn phải mặc đồ bảo hộ.')
    doc.add_paragraph('2. Việc thay thế "Van Oxy Cấp 3" (Oxygen Valve Type-3) là một thao tác cực kỳ rủi ro. Hành động này chỉ được phép thực hiện khi có sự phê duyệt trực tiếp của Kỹ sư trưởng.')
    doc.add_paragraph('3. Để mở khóa kho lưu trữ chứa Van Oxy dự phòng, Kỹ sư trưởng phải nhập Mật mã Override khẩn cấp vào bảng điều khiển trung tâm.')
    doc.save(os.path.join(output_dir, "03_GiaoThuc77A.docx"))

    # 4. PDF: Thông số kỹ thuật
    class PDF(FPDF):
        pass
    pdf = PDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=16)
    pdf.cell(200, 10, text="Bao cao Ky thuat Tau Zeus", new_x="LMARGIN", new_y="NEXT", align='C')
    pdf.set_font("Helvetica", size=12)
    pdf.cell(200, 10, text="Tau Zeus duoc thiet ke voi 2 khoang rieng biet.", new_x="LMARGIN", new_y="NEXT", align='L')
    pdf.cell(200, 10, text="Diem yeu chi mang: He thong Van Oxy Cap 3 (Oxygen Valve Type-3) rat de bi", new_x="LMARGIN", new_y="NEXT", align='L')
    pdf.cell(200, 10, text="hong hoc neu nhiet do ngoai vu tru ha xuong duoi -150 do C.", new_x="LMARGIN", new_y="NEXT", align='L')
    pdf.output(os.path.join(output_dir, "04_BaoCaoZeus.pdf"))

    # 5. CSV: Vật tư
    csv_data = [
        ["Mã Vật Tư", "Tên Linh Kiện", "Số lượng tồn kho", "Vị trí lưu trữ"],
        ["V-01", "Van Oxy Cấp 1", "50", "Kho Tổng Apollo"],
        ["V-02", "Van Oxy Cấp 2", "20", "Kho Hermes"],
        ["V-03", "Van Oxy Cấp 3", "0", "Kho Hermes"],
        ["V-03-Z", "Van Oxy Cấp 3 (Dự phòng khẩn cấp)", "1", "Tủ sắt bị khóa trên tàu Zeus"]
    ]
    with open(os.path.join(output_dir, "05_KhoVatTu.csv"), "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerows(csv_data)

    # 6. PNG: Hình ảnh (Chứa mật mã bị giấu)
    img = Image.new('RGB', (600, 200), color=(50, 50, 50))
    d = ImageDraw.Draw(img)
    # Adding interconnected clues in the image
    d.text((20, 40), "CẢNH BÁO BÃO MẶT TRỜI (SOLAR FLARE) ĐANG DIỄN RA!", fill=(255, 100, 100))
    d.text((20, 80), "Mất kết nối vô tuyến toàn diện giữa tàu Zeus và Hermes.", fill=(255, 255, 255))
    d.text((20, 120), "Mã Override mở tủ sắt khẩn cấp: OMEGA-42", fill=(255, 255, 0))
    img.save(os.path.join(output_dir, "06_BanDoVaMatMa.png"))

    # 7. MP3: Âm thanh khẩn cấp
    audio_text = "Cấp cứu, cấp cứu! Đây là tàu Zeus. Van Oxy cấp 3 của chúng tôi vừa bị vỡ tung. Chúng tôi đang mất dần không khí. Yêu cầu hướng dẫn khẩn cấp!"
    tts = gTTS(audio_text, lang="vi")
    tts.save(os.path.join(output_dir, "07_GhiAmZeus.mp3"))

    print("Successfully generated all complex test files!")

if __name__ == "__main__":
    generate_complex_test_data()
